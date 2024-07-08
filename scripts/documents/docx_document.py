from documents.document import Document
import re
import base64
import sys
import json
from io import BytesIO
from PIL import Image
from docx import Document as _Document
from docx.shared import Inches

class DocxDocument(Document):
    def __init__(self, format, data, template=None):
        super().__init__(format, data, template)
        self.document = _Document()
        self.page_size = {
            "x": 242937,
            "y": 1093491,
            "width": 8656646,
            "height": 3683504
        }
        self.first = True

    def create_page(self):
        if self.first:
            self.first = False
        else:
            self.document.add_page_break()
        return self.document

    def add_title(self, page, title):
        pass

    def add_image(self, page, x, y, width, height, source):
        b64_string = source
        if b64_string.startswith("data:image/"):
            b64_string = b64_string.split(",")[1]
        
        image_data = base64.b64decode(b64_string)
        image_stream = BytesIO(image_data)
        page.add_picture(image_stream, width=Inches(6))

    def add_text(self, page, x, y, width, height, source):
        self.html_to_element(page, source)


